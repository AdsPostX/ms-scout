"""
Attachment ingestion for Scout — Google Sheets URLs + Slack file attachments.

Per-turn input expansion: detect a Sheets URL in @mention text OR extract a file
from event.files[], return a typed result the handler can pass to ask_with_attachment().

Read-only by design. No write-back to external systems.
"""
from __future__ import annotations
import re, base64, ipaddress, subprocess, tempfile, urllib.request, urllib.error, socket
import pathlib
from dataclasses import dataclass
from typing import Literal, Optional
import logging
import pandas as pd
import io

log = logging.getLogger("scout_attachments")

# Tunable constants (single source of truth)
MAX_FILE_BYTES = 10 * 1024 * 1024        # 10MB hard cap
MAX_IMAGE_BYTES = 5 * 1024 * 1024        # 5MB before base64
MAX_TEXT_CHARS = 30_000                   # extracted text cap
PDFTOTEXT_TIMEOUT_S = 10
MAX_REDIRECT_HOPS = 3
ALLOWED_SHEETS_HOSTS = frozenset({"docs.google.com", "accounts.google.com"})
# Google's CSV export 302-redirects to `doc-XX-YY-sheets.googleusercontent.com`
# (where XX-YY are random tokens) to serve the bytes. The `sheets` is part of the
# subdomain string itself, not a sub-label — so a literal `.sheets.googleusercontent.com`
# suffix check would miss it. Tight regex below: optional `[a-z0-9-]+-` prefix
# followed by `sheets.googleusercontent.com` — matches Google's actual pattern,
# rejects prefix-confusion attacks (e.g. `evilsheets.googleusercontent.com`).
_GOOGLE_SHEETS_HOST_RE = re.compile(r'^(?:[a-z0-9-]+-)?sheets\.googleusercontent\.com$')


def _is_sheets_host_allowed(host: str) -> bool:
    """Allow exact hosts plus the doc-*-sheets.googleusercontent.com family."""
    if host in ALLOWED_SHEETS_HOSTS:
        return True
    if _GOOGLE_SHEETS_HOST_RE.match(host):
        return True
    return False

# --- Result type -------------------------------------------------------------

@dataclass
class AttachmentResult:
    kind: Literal["text", "image", "unsupported", "auth_required", "too_large", "error"]
    source: Literal["file", "sheets_url"]
    name: str
    text: Optional[str] = None
    image_b64: Optional[str] = None
    image_media_type: Optional[str] = None
    error: Optional[str] = None

# --- URL detection -----------------------------------------------------------

# Slack delivers links as <url> or <url|label>. Unwrap before matching.
_SLACK_LINK_WRAP_RE = re.compile(r'<(https?://[^|>]+)(?:\|[^>]*)?>')
_SHEETS_URL_RE = re.compile(r'https?://docs\.google\.com/spreadsheets/d/([A-Za-z0-9_-]+)')
_GID_RE = re.compile(r'[?&#]gid=(\d+)')  # Google uses #gid= fragment in normal URLs

def detect_sheets_url(text: str) -> Optional[str]:
    """Return first Google Sheets URL in text, or None. Unwraps Slack's <url|label> first."""
    if not text:
        return None
    unwrapped = _SLACK_LINK_WRAP_RE.sub(r'\1', text)
    m = _SHEETS_URL_RE.search(unwrapped)
    if not m:
        return None
    # Return the full matched URL substring (includes path + query if present)
    return unwrapped[m.start():m.end()] + (
        # Preserve gid if present in the unwrapped text right after the match
        _capture_gid(unwrapped[m.end():])
    )

def _capture_gid(rest: str) -> str:
    """Capture &gid=N or ?gid=N if it appears right after the sheets path."""
    g = _GID_RE.search(rest[:200])  # only look in nearby chars
    return f"&gid={g.group(1)}" if g else ""

# --- Sheets extraction -------------------------------------------------------

def extract_sheets_url(url: str) -> AttachmentResult:
    """Fetch a Sheets URL via export?format=csv, parse as DataFrame, summarize.

    Hard gates: anonymous-only (link-shared), allowlist hosts, block private IPs,
    max 3 redirect hops, login-redirect detection → auth_required.
    """
    name = "google_sheet"
    # Extract sheet ID and optional gid
    m = _SHEETS_URL_RE.search(url)
    if not m:
        return AttachmentResult(kind="error", source="sheets_url", name=name,
                                error="not_a_sheets_url")
    sheet_id = m.group(1)
    g = _GID_RE.search(url)
    export_url = f"https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=csv"
    if g:
        export_url += f"&gid={g.group(1)}"

    try:
        body_bytes, final_url = _safe_fetch(export_url, max_bytes=MAX_FILE_BYTES)
    except _AuthRequired as e:
        return AttachmentResult(kind="auth_required", source="sheets_url", name=name,
                                error=str(e))
    except _TooLarge:
        return AttachmentResult(kind="too_large", source="sheets_url", name=name)
    except Exception as e:
        return AttachmentResult(kind="error", source="sheets_url", name=name,
                                error=f"fetch_failed: {type(e).__name__}: {e}")

    # Detect HTML login page (Sheets returns HTML 200 for unauth, not 401)
    body_head = body_bytes[:200].decode("utf-8", errors="ignore").lstrip()
    if body_head.startswith("<!DOCTYPE html") or body_head.startswith("<html"):
        return AttachmentResult(kind="auth_required", source="sheets_url", name=name,
                                error="login_html_returned")

    try:
        df = pd.read_csv(io.BytesIO(body_bytes))
    except Exception as e:
        return AttachmentResult(kind="error", source="sheets_url", name=name,
                                error=f"csv_parse_failed: {e}")

    summary = _summarize_dataframe(df)
    if len(summary) > MAX_TEXT_CHARS:
        summary = summary[:MAX_TEXT_CHARS] + "\n…[trimmed]"
    return AttachmentResult(kind="text", source="sheets_url", name=name, text=summary)

# --- File extraction ---------------------------------------------------------

# --- Format dispatch table ---------------------------------------------------
# Single source of truth for "what file types can Scout extract?"
# Each entry is (predicate, extractor). Predicates match on lowercase mimetype +
# lowercase extension (including the leading dot, e.g. ".xlsx" or "" if none).
# Extractor signature: (body_bytes, name, mimetype) -> AttachmentResult.
# Order matters — first match wins. To add a new format: implement an extractor
# + add one row. The if/elif chain this replaces grew misaligned (vnd.ms-excel
# was misrouted to read_csv, .xlsx and .docx were silently dropped). Table form
# makes coverage visible and additions one-line.

def _is_pdf(mt: str, ext: str) -> bool:
    return mt == "application/pdf" or ext == ".pdf"

def _is_xlsx(mt: str, ext: str) -> bool:
    # Modern Excel (Office Open XML)
    return (mt == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            or ext == ".xlsx")

def _is_xls(mt: str, ext: str) -> bool:
    # Legacy Excel 97-2003 (binary). Previous code routed this to read_csv which
    # crashes on real .xls bytes — bug, not just gap. Use read_excel + xlrd.
    return mt == "application/vnd.ms-excel" or ext == ".xls"

def _is_docx(mt: str, ext: str) -> bool:
    # Modern Word (Office Open XML). Partner briefs, RFP responses arrive as .docx.
    return (mt == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            or ext == ".docx")

def _is_csv(mt: str, ext: str) -> bool:
    return mt == "text/csv" or ext == ".csv"

def _is_image(mt: str, ext: str) -> bool:
    return (mt in ("image/png", "image/jpeg", "image/gif", "image/webp")
            or (mt.startswith("image/") and ext in (".png", ".jpg", ".jpeg", ".gif", ".webp")))

def _is_text(mt: str, ext: str) -> bool:
    return (mt.startswith("text/") or mt == "application/json"
            or ext in (".txt", ".md", ".markdown", ".json", ".log"))


def _dispatch_image(b: bytes, n: str, m: str) -> "AttachmentResult":
    return _extract_image(b, m, n)


# Order matters — first match wins. Excel/PDF/Docx checked before generic text
# because their mimetypes don't overlap and they have richer parsers.
_EXTRACTORS = [
    (_is_pdf,   lambda b, n, m: _extract_pdf(b, n)),
    (_is_xlsx,  lambda b, n, m: _extract_excel(b, n, engine="openpyxl")),
    (_is_xls,   lambda b, n, m: _extract_excel(b, n, engine="xlrd")),
    (_is_docx,  lambda b, n, m: _extract_docx(b, n)),
    (_is_csv,   lambda b, n, m: _extract_csv(b, n)),
    (_is_image, _dispatch_image),
    (_is_text,  lambda b, n, m: _extract_text(b, n)),
]


def extract_file(file_obj: dict, bot_token: str) -> AttachmentResult:
    """Extract content from a Slack event.files[i] dict.

    Dispatch by the _EXTRACTORS table above: PDF, Excel (xlsx + xls), Word docx,
    CSV, images, plain text/JSON/markdown. Unsupported types degrade gracefully
    with a friendly note; scout_handlers injects the failure into the Claude
    prompt so users get a useful "I couldn't read X, but I can do Y" response.
    """
    name = file_obj.get("name", "unknown")
    mimetype = (file_obj.get("mimetype") or "").lower()
    size = int(file_obj.get("size") or 0)
    url = file_obj.get("url_private", "")
    ext = pathlib.Path(name).suffix.lower()

    if size > MAX_FILE_BYTES:
        return AttachmentResult(kind="too_large", source="file", name=name)

    # Download with Slack bot auth
    try:
        body_bytes = _slack_download(url, bot_token, max_bytes=MAX_FILE_BYTES)
    except Exception as e:
        return AttachmentResult(kind="error", source="file", name=name,
                                error=f"download_failed: {type(e).__name__}: {e}")

    for predicate, extractor in _EXTRACTORS:
        if predicate(mimetype, ext):
            return extractor(body_bytes, name, mimetype)

    return AttachmentResult(kind="unsupported", source="file", name=name,
                            error=f"unsupported_mimetype: {mimetype or 'unknown'} "
                                  f"(extension: {ext or 'none'})")

# --- Internal helpers --------------------------------------------------------

class _AuthRequired(Exception): pass
class _TooLarge(Exception): pass

def _safe_fetch(url: str, max_bytes: int) -> tuple[bytes, str]:
    """Anonymous GET with SSRF guards: host allowlist, no auto-redirect, max hops, IP block."""
    current_url = url
    for hop in range(MAX_REDIRECT_HOPS + 1):
        host = _host_from_url(current_url)
        if not _is_sheets_host_allowed(host):
            raise ValueError(f"host_not_allowed: {host}")
        if _resolves_to_private_ip(host):
            raise ValueError(f"private_ip: {host}")
        req = urllib.request.Request(current_url, headers={"User-Agent": "scout-bot/1.0"})
        opener = urllib.request.build_opener(_NoRedirect())
        try:
            with opener.open(req, timeout=15) as resp:
                # If we get here, no redirect occurred
                if int(resp.headers.get("Content-Length") or 0) > max_bytes:
                    raise _TooLarge()
                data = resp.read(max_bytes + 1)
                if len(data) > max_bytes:
                    raise _TooLarge()
                return data, current_url
        except urllib.error.HTTPError as e:
            if e.code in (301, 302, 303, 307, 308):
                next_url = e.headers.get("Location", "")
                if not next_url:
                    raise ValueError("redirect_no_location")
                # Detect login redirect specifically
                if "accounts.google.com" in next_url:
                    raise _AuthRequired("redirected_to_login")
                current_url = next_url
                continue
            if e.code in (401, 403):
                raise _AuthRequired(f"http_{e.code}")
            raise
    raise ValueError(f"too_many_redirects: {MAX_REDIRECT_HOPS}")

class _NoRedirect(urllib.request.HTTPRedirectHandler):
    """Block urllib's auto-redirect so we can validate each hop."""
    def redirect_request(self, req, fp, code, msg, headers, newurl):
        return None

def _host_from_url(url: str) -> str:
    from urllib.parse import urlparse
    return (urlparse(url).hostname or "").lower()

def _resolves_to_private_ip(host: str) -> bool:
    """Check ALL A/AAAA records (not just first) so a host with both public and
    private IPs is rejected. Uses getaddrinfo to enumerate every address family.
    Fail-closed on resolution error — if we can't resolve, don't fetch."""
    try:
        infos = socket.getaddrinfo(host, None, type=socket.SOCK_STREAM)
    except Exception:
        return True  # fail-closed
    for family, _type, _proto, _canon, sockaddr in infos:
        ip_str = sockaddr[0]
        try:
            addr = ipaddress.ip_address(ip_str)
        except ValueError:
            return True  # unparseable IP — fail-closed
        if addr.is_private or addr.is_loopback or addr.is_link_local or addr.is_reserved or addr.is_multicast:
            return True
    return False

def _slack_download(url: str, bot_token: str, max_bytes: int) -> bytes:
    """Download a Slack file via url_private with bot token. Slack CDN only — cap size."""
    # url_private always points at files.slack.com (the Slack CDN). Narrower
    # than the original `slack.com` prefix to reduce attack surface.
    if not url.startswith("https://files.slack.com/"):
        raise ValueError(f"unexpected_slack_url: {url[:80]}")
    req = urllib.request.Request(url, headers={"Authorization": f"Bearer {bot_token}"})
    with urllib.request.urlopen(req, timeout=15) as resp:
        data = resp.read(max_bytes + 1)
        if len(data) > max_bytes:
            raise _TooLarge()
        return data

def _extract_pdf(body_bytes: bytes, name: str) -> AttachmentResult:
    """Run pdftotext via subprocess on a tempfile. No shell, explicit list args, timeout."""
    fd, path = tempfile.mkstemp(suffix=".pdf")
    try:
        import os as _os
        with _os.fdopen(fd, "wb") as f:
            f.write(body_bytes)
        try:
            result = subprocess.run(
                ["pdftotext", "-layout", path, "-"],
                capture_output=True, timeout=PDFTOTEXT_TIMEOUT_S, check=False,
            )
        except subprocess.TimeoutExpired:
            return AttachmentResult(kind="error", source="file", name=name,
                                    error="pdf_parse_timeout")
        except FileNotFoundError:
            # pdftotext not installed — try pdfplumber fallback
            return _extract_pdf_fallback(body_bytes, name)
        if result.returncode != 0:
            return AttachmentResult(kind="error", source="file", name=name,
                                    error=f"pdftotext_exit_{result.returncode}")
        text = result.stdout.decode("utf-8", errors="ignore")
        if len(text) > MAX_TEXT_CHARS:
            text = text[:MAX_TEXT_CHARS] + "\n…[trimmed]"
        return AttachmentResult(kind="text", source="file", name=name, text=text)
    finally:
        try:
            import os as _os
            _os.unlink(path)
        except Exception:
            pass

def _extract_pdf_fallback(body_bytes: bytes, name: str) -> AttachmentResult:
    """pdfplumber fallback if pdftotext CLI is missing."""
    try:
        import pdfplumber
    except ImportError:
        return AttachmentResult(kind="error", source="file", name=name,
                                error="no_pdf_extractor_available")
    try:
        with pdfplumber.open(io.BytesIO(body_bytes)) as pdf:
            text = "\n".join((page.extract_text() or "") for page in pdf.pages)
        if len(text) > MAX_TEXT_CHARS:
            text = text[:MAX_TEXT_CHARS] + "\n…[trimmed]"
        return AttachmentResult(kind="text", source="file", name=name, text=text)
    except Exception as e:
        return AttachmentResult(kind="error", source="file", name=name,
                                error=f"pdfplumber_failed: {e}")

def _extract_csv(body_bytes: bytes, name: str) -> AttachmentResult:
    try:
        df = pd.read_csv(io.BytesIO(body_bytes))
    except Exception as e:
        return AttachmentResult(kind="error", source="file", name=name,
                                error=f"csv_parse_failed: {e}")
    summary = _summarize_dataframe(df)
    if len(summary) > MAX_TEXT_CHARS:
        summary = summary[:MAX_TEXT_CHARS] + "\n…[trimmed]"
    return AttachmentResult(kind="text", source="file", name=name, text=summary)


def _extract_excel(body_bytes: bytes, name: str, engine: str = "openpyxl") -> AttachmentResult:
    """Extract content from .xlsx (engine='openpyxl') or .xls (engine='xlrd').

    Reads ALL sheets but summarizes the first one only — keeps the response
    bounded. Sheet names are surfaced in the summary so users can ask Scout
    to drill into a specific sheet in a follow-up turn.

    Caught by Sidd's live testing on PR #236: AT&T Views .xlsx fell through
    to "unsupported" because the dispatch only knew CSV/PDF/image/text.
    """
    # Lazy enumeration via pd.ExcelFile — `sheet_name=None` on read_excel would
    # eagerly materialize every sheet into memory even though we only summarize
    # the first one (caught by CodeRabbit on PR #238). For a workbook with many
    # tabs the difference is meaningful in memory + parse latency.
    try:
        with pd.ExcelFile(io.BytesIO(body_bytes), engine=engine) as xls:
            sheet_names = list(xls.sheet_names)
            if not sheet_names:
                return AttachmentResult(kind="error", source="file", name=name,
                                        error="excel_empty_or_no_sheets")
            first_name = sheet_names[0]
            first_df = pd.read_excel(xls, sheet_name=first_name)
    except ImportError as e:
        return AttachmentResult(kind="error", source="file", name=name,
                                error=f"excel_engine_missing: install {engine}: {e}")
    except Exception as e:
        return AttachmentResult(kind="error", source="file", name=name,
                                error=f"excel_parse_failed: {type(e).__name__}: {e}")

    if len(sheet_names) > 1:
        header = (
            f"Workbook has {len(sheet_names)} sheets: {sheet_names!r}. "
            f"Analyzing first sheet ({first_name!r}); ask to drill into others by name.\n"
        )
    else:
        header = f"Sheet: {first_name!r}\n"

    summary = header + _summarize_dataframe(first_df)
    if len(summary) > MAX_TEXT_CHARS:
        summary = summary[:MAX_TEXT_CHARS] + "\n…[trimmed]"
    return AttachmentResult(kind="text", source="file", name=name, text=summary)


def _extract_docx(body_bytes: bytes, name: str) -> AttachmentResult:
    """Extract paragraphs + table cells from a .docx (Office Open XML) file.

    python-docx skips embedded images; we get text only. Tables become
    tab-separated rows so Claude can still read tabular content. Real use
    case: partner briefs, RFP responses, integration specs in ad-tech all
    arrive as .docx, not text.
    """
    try:
        from docx import Document  # python-docx
    except ImportError:
        return AttachmentResult(kind="error", source="file", name=name,
                                error="docx_engine_missing: install python-docx")
    try:
        doc = Document(io.BytesIO(body_bytes))
    except Exception as e:
        return AttachmentResult(kind="error", source="file", name=name,
                                error=f"docx_parse_failed: {type(e).__name__}: {e}")

    parts = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append("\t".join(cells))

    text = "\n".join(parts) if parts else "(empty document)"
    if len(text) > MAX_TEXT_CHARS:
        text = text[:MAX_TEXT_CHARS] + "\n…[trimmed]"
    return AttachmentResult(kind="text", source="file", name=name, text=text)


def _extract_image(body_bytes: bytes, mimetype: str, name: str) -> AttachmentResult:
    if len(body_bytes) > MAX_IMAGE_BYTES:
        return AttachmentResult(kind="too_large", source="file", name=name)
    b64 = base64.b64encode(body_bytes).decode("ascii")
    return AttachmentResult(kind="image", source="file", name=name,
                            image_b64=b64, image_media_type=mimetype)

def _extract_text(body_bytes: bytes, name: str) -> AttachmentResult:
    try:
        text = body_bytes.decode("utf-8", errors="replace")
    except Exception as e:
        return AttachmentResult(kind="error", source="file", name=name,
                                error=f"text_decode_failed: {e}")
    if len(text) > MAX_TEXT_CHARS:
        text = text[:MAX_TEXT_CHARS] + "\n…[trimmed]"
    return AttachmentResult(kind="text", source="file", name=name, text=text)

def _summarize_dataframe(df: "pd.DataFrame") -> str:
    """Compact markdown summary for any tabular source. Shared by CSV file + Sheets URL."""
    rows, cols = df.shape
    head_md = df.head(10).to_markdown(index=False)
    # Numeric describe, if any numeric columns exist
    try:
        numeric_df = df.select_dtypes(include="number")
        stats_md = numeric_df.describe().to_markdown() if not numeric_df.empty else "(no numeric columns)"
    except Exception:
        stats_md = "(describe failed)"
    return (
        f"Shape: {rows} rows × {cols} columns\n"
        f"Columns: {list(df.columns)}\n\n"
        f"Head (first 10 rows):\n{head_md}\n\n"
        f"Numeric stats:\n{stats_md}"
    )
